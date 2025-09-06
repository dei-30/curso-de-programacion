import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from gtts import gTTS
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from docx2pdf import convert
from PyPDF2 import PdfMerger
import os

def requiere_revision(funcion_bono):
    """Decorador que marca a un empleado para revisión si su bono es 0."""
    def wrapper(self):
        bono = funcion_bono(self)
        if bono == 0:
            print(f"-> REVISIÓN REQUERIDA para {self.nombre}. Bono calculado: 0.")
            self.requiere_revision = True
        return bono
    return wrapper

# --- Clases de Empleados ---
class Empleado:
    """Clase base para un empleado."""
    def __init__(self, id_empleado, nombre, email, salario_base, kpis):
        self.id_empleado = id_empleado
        self.nombre = nombre
        self.email = email
        self.__salario = salario_base  # Encapsulamiento
        self.kpis = kpis
        self.requiere_revision = False

    def get_salario(self):
        return self.__salario

    @requiere_revision
    def calcular_bono(self):
        """Calcula el bono basado en el promedio de KPIs. Polimorfismo."""
        promedio_kpi = sum(self.kpis.values()) / len(self.kpis)
        if promedio_kpi > 3.5:
            return self.__salario * 0.10
        return 0
    
class EmpleadoComercial(Empleado):
    """Clase para empleados del área comercial con un cálculo de bono diferente."""
    def __init__(self, id_empleado, nombre, email, salario_base, kpis, ventas_totales):
        super().__init__(id_empleado, nombre, email, salario_base, kpis)
        self.ventas_totales = ventas_totales

    @requiere_revision
    def calcular_bono(self):
        """Cálculo de bono polimórfico para comerciales."""
        bono_kpi = super().calcular_bono()
        bono_ventas = self.ventas_totales * 0.02
        # El bono es el mayor de los dos, no la suma, para no duplicar.
        return max(bono_kpi, bono_ventas)
    
# --- Funciones del Sistema ---
def leer_datos_empleados(ruta_excel):
    """Lee los datos y KPIs de los empleados desde un archivo Excel."""
    try:
        df = pd.read_excel(ruta_excel)
        empleados = []
        for _, row in df.iterrows():
            kpis = {
                "Comunicación": row["Comunicación"],
                "Trabajo en Equipo": row["Trabajo en Equipo"],
                "Iniciativa": row["Iniciativa"],
                "Productividad": row["Productividad"],
                "Adaptabilidad": row["Adaptabilidad"]
            }
            if row["Departamento"] == "Comercial":
                empleado = EmpleadoComercial(
                    row["ID"], row["Nombre"], row["Email"], row["Salario"], kpis, row["Ventas"]
                )
            else:
                empleado = Empleado(
                    row["ID"], row["Nombre"], row["Email"], row["Salario"], kpis
                )
            empleados.append(empleado)
        return empleados
    except FileNotFoundError:
        print(f"Error: No se encontró el archivo {ruta_excel}")
        return []
    except Exception as e:
        print(f"Error al leer el Excel: {e}")
        return []
    
def generar_grafico_radar(empleado, ruta_salida):
    """Genera y guarda un gráfico de radar para los KPIs de un empleado."""
    labels = list(empleado.kpis.keys())
    stats = list(empleado.kpis.values())
    
    angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
    stats += stats[:1]
    angles += angles[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
    ax.fill(angles, stats, color='red', alpha=0.25)
    ax.plot(angles, stats, color='red', linewidth=2)
    
    ax.set_yticklabels([])
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(labels)
    ax.set_title(f"Desempeño de {empleado.nombre}", size='large', color='blue', y=1.1)
    
    ruta_grafico = os.path.join(ruta_salida, f"grafico_{empleado.id_empleado}.png")
    plt.savefig(ruta_grafico)
    plt.close(fig)
    return ruta_grafico

def generar_informe_word(empleado, bono, ruta_grafico, ruta_salida):
    """Genera un informe de desempeño personalizado en formato Word."""
    doc = Document()
    doc.add_heading(f"Informe de Desempeño - {empleado.nombre}", 0)
    
    doc.add_paragraph(f"ID de Empleado: {empleado.id_empleado}")
    doc.add_paragraph(f"Salario Base: ${empleado.get_salario():,.2f}")
    
    doc.add_heading("Evaluación de KPIs", level=1)
    if os.path.exists(ruta_grafico):
        doc.add_picture(ruta_grafico, width=Inches(5.0))
    
    doc.add_heading("Resumen y Bonificación", level=1)
    promedio_kpi = sum(empleado.kpis.values()) / len(empleado.kpis)
    doc.add_paragraph(f"El promedio de KPIs de {empleado.nombre} es de {promedio_kpi:.2f} sobre 5.")
    doc.add_paragraph(f"Bono calculado para este período: ${bono:,.2f}")
    
    if empleado.requiere_revision:
        doc.add_paragraph("\n**NOTA: Este empleado requiere una revisión de desempeño por parte de RRHH.**").bold = True
        
    ruta_documento = os.path.join(ruta_salida, f"informe_{empleado.id_empleado}.docx")
    doc.save(ruta_documento)
    return ruta_documento

def generar_audio_resumen(empleado, bono, ruta_salida):
    """Genera un archivo de audio con el resumen del informe."""
    promedio_kpi = sum(empleado.kpis.values()) / len(empleado.kpis)
    texto = f"Resumen de desempeño para {empleado.nombre}. Promedio de KPIs: {promedio_kpi:.2f}. Bono asignado: ${bono:,.2f}."
    if empleado.requiere_revision:
        texto += " Se requiere revisión adicional por parte de Recursos Humanos."
        
    tts = gTTS(text=texto, lang='es')
    ruta_audio = os.path.join(ruta_salida, f"resumen_{empleado.id_empleado}.mp3")
    tts.save(ruta_audio)
    return ruta_audio

def enviar_correo(destinatario, asunto, cuerpo, adjuntos, config_email, nombre_empleado_log):
    """Envía un correo electrónico con su informe."""
    if not config_email or config_email["usuario"] == "tu_email@gmail.com":
        print(f"ADVERTENCIA: Correo para {nombre_empleado_log} no enviado (configuración de email pendiente).")
        return

    msg = MIMEMultipart()
    msg['From'] = config_email["usuario"]
    msg['To'] = destinatario
    msg['Subject'] = asunto
    msg.attach(MIMEText(cuerpo, 'plain'))

    for ruta_adjunto in adjuntos:
        if os.path.exists(ruta_adjunto):
            part = MIMEBase('application', 'octet-stream')
            with open(ruta_adjunto, 'rb') as archivo:
                part.set_payload(archivo.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f"attachment; filename={os.path.basename(ruta_adjunto)}")
            msg.attach(part)
    
    try:
        with smtplib.SMTP(config_email["smtp_server"], config_email["puerto"]) as server:
            server.starttls()
            server.login(config_email["usuario"], config_email["contrasena"])
            server.send_message(msg)
        print(f"Correo enviado exitosamente a {destinatario}.")
    except Exception as e:
        print(f"Error al enviar correo a {destinatario}: {e}")

def consolidar_pdfs(rutas_docs, ruta_salida_final):
    """Convierte los Word a PDF y los une en un solo PDF maestro."""
    pdfs_generados = []
    print("\nIniciando conversión de Word a PDF...")
    for doc_path in rutas_docs:
        try:
            convert(doc_path)
            pdf_path = doc_path.replace(".docx", ".pdf")
            if os.path.exists(pdf_path):
                pdfs_generados.append(pdf_path)
                print(f"  - Convertido: {os.path.basename(doc_path)}")
        except Exception as e:
            print(f"Error convirtiendo {doc_path}: {e}")

    if not pdfs_generados:
        print("No se generaron PDFs para consolidar.")
        return

    print("\nIniciando consolidación de PDFs...")
    merger = PdfMerger()
    for pdf in pdfs_generados:
        merger.append(pdf)
    
    merger.write(ruta_salida_final)
    merger.close()
    print(f"PDF maestro creado en: {ruta_salida_final}")

    # Opcional: limpiar los PDFs individuales
    for pdf in pdfs_generados:
        os.remove(pdf)

# --- Ejecución Principal ---
if __name__ == "__main__":
    # --- Crear directorios de salida ---
    if not os.path.exists("informes_individuales"):
        os.makedirs("informes_individuales")
    
    # --- Crear Excel de ejemplo ---
    datos_ejemplo = {
        "ID": [101, 102, 103, 104],
        "Nombre": ["Ana Torres", "Carlos Vera", "Lucía Mora", "Pedro Nieto"],
        "Email": ["ana.t@example.com", "carlos.v@example.com", "lucia.m@example.com", "pedro.n@example.com"],
        "Departamento": ["TI", "Comercial", "RRHH", "Comercial"],
        "Salario": [50000, 55000, 48000, 52000],
        "Ventas": [0, 150000, 0, 80000],
        "Comunicación": [4, 5, 3, 2],
        "Trabajo en Equipo": [5, 4, 4, 3],
        "Iniciativa": [4, 5, 3, 1],
        "Productividad": [3, 4, 5, 2],
        "Adaptabilidad": [4, 3, 4, 3]
    }
    df_ejemplo = pd.DataFrame(datos_ejemplo)
    ruta_excel = "datos_empleados.xlsx"
    df_ejemplo.to_excel(ruta_excel, index=False)

    # --- Configuración de Email (MODIFICAR) ---
    config_email = {
        "smtp_server": "smtp.gmail.com",
        "puerto": 587,
        "usuario": "zambranoeli2010@gmail.com",
        "contrasena": "hcqmwdurcnlzluek" 
    }
    # --- Destinatario Adicional (MODIFICAR) ---
    destinatario_profesor = "deivisfonseca3006@gmail.com" # <-- Poner el email del profe aquí. Dejar en blanco "" para no enviar.

    # --- Procesar Empleados ---
    print("Iniciando sistema de informes de RRHH...")
    empleados = leer_datos_empleados(ruta_excel)
    rutas_docs_generados = []

    for emp in empleados:
        print(f"\nProcesando a: {emp.nombre}")
        
        # 1. Calcular bono
        bono = emp.calcular_bono()
        
        # 2. Generar gráfico
        ruta_grafico = generar_grafico_radar(emp, "informes_individuales")

        # 3. Generar Word
        ruta_word = generar_informe_word(emp, bono, ruta_grafico, "informes_individuales")
        rutas_docs_generados.append(ruta_word)
        
        # 4. Generar Audio
        ruta_audio = generar_audio_resumen(emp, bono, "informes_individuales")
        
        # 5. Enviar Email
        # A. Enviar al empleado
        asunto_empleado = f"Tu Informe de Desempeño - {emp.nombre}"
        cuerpo_empleado = f"Hola {emp.nombre.split()[0]},\n\nAdjunto encontrarás tu informe de desempeño y un resumen en audio.\n\nSaludos,\nEquipo de RRHH."
        enviar_correo(emp.email, asunto_empleado, cuerpo_empleado, [ruta_word, ruta_audio], config_email, emp.nombre)

    # B. Enviar copia al profe(si está configurado)
        if destinatario_profesor:
            asunto_profesor = f"Copia Informe: {emp.nombre}"
            cuerpo_profesor = f"Copia del informe de desempeño para el empleado {emp.nombre} (ID: {emp.id_empleado})."
            enviar_correo(destinatario_profesor, asunto_profesor, cuerpo_profesor, [ruta_word, ruta_audio], config_email, emp.nombre)

    # 6. Consolidar todos los informes en un PDF maestro
    if rutas_docs_generados:
        consolidar_pdfs(rutas_docs_generados, "informe_maestro_consolidado.pdf")

    print("\nProceso completado.")
