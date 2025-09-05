# --- Ejercicio 1: Exploración y Carga ---
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Cargar el archivo
try:
    df_ventas = pd.read_excel('ventas_tienda.3.xlsx')
    print("DataFrame cargado con éxito. Exploración inicial:")
    print(df_ventas.head(10))
    print("\n--- Información general del DataFrame ---")
    df_ventas.info()
    print("\n--- Estadísticas descriptivas ---")
    print(df_ventas.describe())
except FileNotFoundError:
    print("Error: El archivo 'ventas_tienda.3.xlsx' no se encontró. Asegúrate de que está en la misma carpeta.")
    exit()

df_ventas['Cliente'].fillna('Consumidor Final', inplace=True)

# Convertir la columna 'Fecha' a tipo datetime
df_ventas['Fecha'] = pd.to_datetime(df_ventas['Fecha'])

# Crear la nueva columna 'Ingreso_Total'
df_ventas['Ingreso_Total'] = df_ventas['Unidades_Vendidas'] * df_ventas['Precio_Unitario']

print("\n--- DataFrame después de la limpieza y creación de columna ---")
print(df_ventas.head())

# --- Ejercicio 3: Agrupación, Visualización y Exportación ---
# Agrupar por 'Categoría' y sumar 'Ingreso_Total'
ingresos_por_categoria = df_ventas.groupby('Categoría')['Ingreso_Total'].sum().reset_index()

# Ordenar los resultados
ingresos_por_categoria_ordenado = ingresos_por_categoria.sort_values(by='Ingreso_Total', ascending=False)

print("\n--- Ingresos totales por categoría (ordenado) ---")
print(ingresos_por_categoria_ordenado)

# Generar el gráfico de barras
plt.figure(figsize=(10, 6))
sns.barplot(
    x='Categoría',
    y='Ingreso_Total',
    data=ingresos_por_categoria_ordenado,
    palette='viridis'
)
plt.title('Ingresos Totales por Categoría de Producto', fontsize=16)
plt.xlabel('Categoría', fontsize=12)
plt.ylabel('Ingreso Total ($)', fontsize=12)

# Añadir etiquetas de valor
for index, row in ingresos_por_categoria_ordenado.iterrows():
    plt.text(index, row['Ingreso_Total'] + 50, f'${row["Ingreso_Total"]:.2f}', ha='center')

plt.show()

# Exportar el DataFrame agrupado a un nuevo Excel
ingresos_por_categoria_ordenado.to_excel('reporte_categorias.xlsx', index=False)
print("\n¡Análisis completado! El reporte se ha exportado a 'reporte_categorias.xlsx'.")


df_ventas['Fecha'] = pd.to_datetime(df_ventas['Fecha'])
df_ventas['Mes'] = df_ventas['Fecha'].dt.month

# Agrupar por mes y sumar el ingreso total
ingresos_mensuales = df_ventas.groupby('Mes')['Ingreso_Total'].sum().reset_index()

print("Ingresos mensuales:")
print(ingresos_mensuales)

# Crear el gráfico de líneas de tendencia
plt.figure(figsize=(10, 6))
plt.plot(ingresos_mensuales['Mes'], ingresos_mensuales['Ingreso_Total'], marker='o', linestyle='-', color='b')
plt.title('Tendencia de Ingresos Mensuales', fontsize=16)
plt.xlabel('Mes', fontsize=12)
plt.ylabel('Ingreso Total ($)', fontsize=12)
plt.xticks(ingresos_mensuales['Mes'])
plt.grid(True)

# Guardar el gráfico como un archivo PNG
plt.savefig('tendencia_mensual.png')
print("\nGráfico de tendencia mensual guardado como 'tendencia_mensual.png'.")

# Encontrar el mes con el mayor ingreso para el reporte
mes_estrella = ingresos_mensuales.loc[ingresos_mensuales['Ingreso_Total'].idxmax()]
mes_mayor_venta = int(mes_estrella['Mes'])
ingreso_mayor_venta = mes_estrella['Ingreso_Total']

# Crear un nuevo documento de Word
doc = Document()

# Agregar un título al documento
doc.add_heading('Reporte de Tendencia de Ventas Mensuales', level=1)

# Añadir un párrafo resumiendo el mes de mayor venta
parrafo = doc.add_paragraph(f"El análisis de ventas muestra que el mes con el mayor volumen de ventas fue el mes {mes_mayor_venta}, generando un total de ${ingreso_mayor_venta:.2f}. Esto destaca una tendencia positiva en ese período, como se puede observar en el gráfico a continuación.")

# Insertar el gráfico de líneas en el documento
doc.add_picture('tendencia_mensual.png', width=Inches(6))

# Guardar el documento final
doc.save('reporte_ventas_mensuales.docx')
print("\n¡Reporte de ventas en Word creado y guardado como 'reporte_ventas_mensuales.docx'!")

# Exportar el DataFrame completo y limpio a un nuevo archivo de Excel
df_ventas.to_excel('datos_procesados_final.xlsx', index=False)
print("DataFrame completo exportado a 'datos_procesados_final.xlsx'.")


doc = Document()

# 2. Título e introducción
doc.add_heading('Informe de Análisis de Ventas Trimestral', level=1)
doc.add_paragraph('Este informe presenta un análisis detallado de las ventas de la tienda, destacando el rendimiento de las categorías de productos y las tendencias de ingresos mensuales. El objetivo es ofrecer una visión clara para la toma de decisiones estratégicas.')

# 3. Insertar la tabla de ingresos por categoría
doc.add_heading('Rendimiento por Categoría de Producto', level=2)
# Convertir el DataFrame a una tabla en Word
table = doc.add_table(ingresos_por_categoria_ordenado.shape[0] + 1, ingresos_por_categoria_ordenado.shape[1])
table.style = 'Table Grid'

# Añadir encabezados de la tabla
for j in range(ingresos_por_categoria_ordenado.shape[-1]):
    table.cell(0, j).text = ingresos_por_categoria_ordenado.columns[j]

# Añadir datos de la tabla
for i in range(ingresos_por_categoria_ordenado.shape[0]):
    for j in range(ingresos_por_categoria_ordenado.shape[-1]):
        table.cell(i + 1, j).text = str(ingresos_por_categoria_ordenado.values[i, j])

# 4. Insertar el gráfico de tendencias mensuales
doc.add_heading('Tendencia de Ingresos Mensuales', level=2)
doc.add_paragraph('La siguiente gráfica muestra el comportamiento de las ventas a lo largo de los meses, identificando el mes de mayor rendimiento.')
doc.add_picture('tendencia_mensual.png', width=Inches(6))

# 5. Añadir un párrafo de conclusiones
doc.add_heading('Conclusiones Clave', level=2)
doc.add_paragraph(f"El análisis confirma que la categoría de productos '{ingresos_por_categoria_ordenado.iloc[0]['Categoría']}' es la principal fuente de ingresos. Además, la tendencia de ventas muestra que el mes {int(ingresos_mensuales.loc[ingresos_mensuales['Ingreso_Total'].idxmax()]['Mes'])} fue el más exitoso, lo que sugiere un posible pico estacional. Se recomienda investigar las estrategias de marketing o los factores externos que contribuyeron a este pico para replicar el éxito en futuros períodos.")

# 6. Guardar el informe final
doc.save('reporte_final_consolidado.docx')
print("\n¡El informe maestro en Word ha sido creado y guardado como 'reporte_final_consolidado.docx'!")