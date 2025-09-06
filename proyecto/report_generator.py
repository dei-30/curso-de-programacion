from abc import ABC, abstractmethod
import os
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from gtts import gTTS
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from fpdf import FPDF
import time

def log_operacion(cls):
    def accion(self, *args, **kwargs):
        log_entry = f"[{time.strftime('%Y-%m-%d %H:%M:%S')}] Ejecutando: {cls.__name__}..."
        self.log.append(log_entry)
        print(log_entry)
        result = cls(self, *args, **kwargs)
        self.log.append(f"[{time.strftime('%Y-%m-%d %H:%M:%S')}] {cls.__name__} completado.")
        return result
    return accion   


class GeneradorReporte(ABC):
    def __init__(self, excel_path, output_dir):
        self.excel_path = excel_path
        self.output_dir = output_dir
        self.log = []

    @abstractmethod
    def generar_reporte(self):
        pass


class ReporteVentas(GeneradorReporte):
    def __init__(self, excel_path, output_dir):
        super().__init__(excel_path, output_dir)
        self.__dataframe = None
        self.resumen_ventas = {}
        self.grafico_path = ""
        self.word_path = ""
        self.audio_path = ""
        self.log_path = ""

    @log_operacion
    def _cargar_datos(self):
        try:
            self.__dataframe = pd.read_excel(self.excel_path)
            print("✔️ Datos cargados con éxito.")
        except Exception as e:
            self.log.append(f"❌ Error al cargar datos: {e}")
            FileNotFoundError(f"Error al cargar datos del archivo: {self.excel_path}")

    @log_operacion
    def _analizar_ventas(self):
        if self.__dataframe is None:
            ValueError("DataFrame no cargado. Por favor, cargue los datos primero.")
        
        ventas_totales = self.__dataframe['Ventas'].sum()
        producto_mas_vendido = self.__dataframe.groupby('Producto')['Ventas'].sum().idxmax()
        ventas_por_region = self.__dataframe.groupby('Región')['Ventas'].sum()
        
        self.resumen_ventas = {
            'ventas_totales': ventas_totales,
            'producto_mas_vendido': producto_mas_vendido,
            'ventas_por_region': ventas_por_region
        }
        print("✔️ Análisis de ventas completado.")

    @log_operacion
    def _crear_grafico(self):
        plt.style.use('ggplot')
        ventas_por_region = self.resumen_ventas['ventas_por_region']
        plt.figure(figsize=(10, 6))
        ventas_por_region.plot(kind='bar', color='skyblue')
        plt.title('Ventas Totales por Región')
        plt.xlabel('Región')
        plt.ylabel('Ventas ($)')
        plt.xticks(rotation=45)
        plt.tight_layout()
        self.grafico_path = os.path.join(self.output_dir, "ventas_por_region.png")
        plt.savefig(self.grafico_path)
        plt.close()
        print(f"✔️ Gráfico guardado en {self.grafico_path}.")

    @log_operacion
    def _crear_word(self):
        document = Document()
        document.add_heading('Reporte de Ventas Automatizado', 0)
        document.add_heading('Resumen General', level=1)
        document.add_paragraph(f"Ventas totales: ${self.resumen_ventas['ventas_totales']:.2f}")
        document.add_paragraph(f"Producto más vendido: {self.resumen_ventas['producto_mas_vendido']}")
        document.add_heading('Ventas por Región', level=1)
        document.add_picture(self.grafico_path, width=Inches(6))
        self.word_path = os.path.join(self.output_dir, "Reporte_Ventas.docx")
        document.save(self.word_path)
        print(f"✔️ Reporte en Word guardado en {self.word_path}.")

    @log_operacion
    def _generar_audio(self):
        texto_resumen = (
            f"El reporte de ventas ha sido generado. Las ventas totales ascienden a "
            f"${self.resumen_ventas['ventas_totales']:.2f}. El producto más vendido es "
            f"{self.resumen_ventas['producto_mas_vendido']}."
        )
        tts = gTTS(text=texto_resumen, lang='es', slow=False)
        self.audio_path = os.path.join(self.output_dir, "resumen_ventas.mp3")
        tts.save(self.audio_path)
        print(f"✔️ Audio de resumen guardado en {self.audio_path}.")

    @log_operacion
    def _enviar_correo(self, rem, passw, dest):
        msg = MIMEMultipart()
        msg['From'] = rem
        msg['To'] = dest
        msg['Subject'] = "Reporte de Ventas Automático"
        
        body = "Adjunto encontrarás el reporte de ventas, el gráfico y el resumen en audio."
        msg.attach(MIMEText(body, 'plain'))
        
        for file_path in [self.word_path, self.grafico_path, self.audio_path]:
            try:
                part = MIMEBase('application', 'octet-stream')
                with open(file_path, 'rb') as f:
                    part.set_payload(f.read())
                encoders.encode_base64(part)
                part.add_header('Content-Disposition', f"attachment; filename= {file_path.split('/')[-1]}")
                msg.attach(part)
            except FileNotFoundError:
                print(f"❌ Archivo no encontrado: {file_path}")
        
        try:
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(rem, passw)
            text = msg.as_string()
            server.sendmail(rem, dest, text)
            server.quit()
            print("✔️ Correo enviado con éxito.")
        except Exception as e:
            self.log.append(f"❌ Error al enviar correo: {e}")
            raise ConnectionError(f"Error al enviar correo: {e}")

    @log_operacion
    def _generar_pdf_log(self):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Log del Proceso de Reporte de Ventas", ln=True, align='C')
        for entry in self.log:
            pdf.cell(200, 10, txt=entry, ln=True)
        self.log_path = os.path.join(self.output_dir, "proceso_log.pdf")
        pdf.output(self.log_path)
        print(f"✔️ Log en PDF guardado en {self.log_path}.")

    def generar_reporte(self):
        self._cargar_datos()
        self._analizar_ventas()
        self._crear_grafico()
        self._crear_word()
        self._generar_audio()
        self._generar_pdf_log()